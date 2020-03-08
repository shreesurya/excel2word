# -*- coding: utf-8 -*-
"""
Created on Fri Dec 27 10:48:32 2019

@author: shrie
"""

# Program extracting first column 
import xlrd
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

lst = []
loc = ("D:/Seva Project Tag Making/Jan_HP_Tags.xlsx") 
document = Document()
#paragraph = document.add_paragraph()
#paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
style = document.styles['Normal']
font = style.font
font.name = 'Algerian'
font.size = Pt(40)
#style.alignment = 'Center'
wb = xlrd.open_workbook(loc) 
sheet = wb.sheet_by_index(0) 
#sheet.cell_value(0, 0) 

for i in range(sheet.nrows): 
    #print(sheet.cell_value(i, 1))
    name = sheet.cell_value(i, 1)
    print(name.upper())
    lst.append(name.upper())
    #a = ("b", "g", "a", "d", "f", "c", "h", "e")
x = sorted(lst)
table = document.add_table(rows=sheet.nrows, cols=2)
#table.alignment = 'Center'
hdr_cells = table.rows[0].cells
#hdr_cells[0].text = 'Id'
#hdr_cells[1].text = 'Quantity'
#hdr_cells[2].text = 'Description'

for j in range(sheet.nrows):
    hdr_cells = table.rows[j].cells
    hdr_cells[0].text = lst[j]
    hdr_cells[1].text = lst[j]
    

document.save('simple.docx')